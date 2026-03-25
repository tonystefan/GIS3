"""
Docx service: gera arquivo .docx formatado da ata usando python-docx.
"""
import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

from apps.transcricao.models import Ata


def _set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula de tabela."""
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_color}" w:color="auto" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


class DocxService:

    @staticmethod
    def gerar_ata_docx(ata: Ata) -> bytes:
        """Gera o documento .docx da ata e retorna os bytes."""
        doc = Document()

        # Configurar margens
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Estilo padrão
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # --- Cabeçalho ---
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run('ATA DE REUNIÃO')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitulo.add_run(ata.titulo)
        run2.font.size = Pt(13)
        run2.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        doc.add_paragraph()

        # --- Informações gerais ---
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'

        labels = ['Data', 'Local', 'Participantes']
        values = [
            ata.data_reuniao.strftime('%d/%m/%Y') if isinstance(ata.data_reuniao, date) else str(ata.data_reuniao),
            ata.local or '—',
            ata.participantes or '—',
        ]

        # Terceira linha: mesclar células para participantes
        for i, (label, value) in enumerate(zip(labels, values)):
            row = info_table.rows[i]
            cell_label = row.cells[0]
            cell_value = row.cells[1]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].bold = True
            cell_value.text = value
            _set_cell_bg(cell_label, 'EEF2FF')

        doc.add_paragraph()

        # --- Seções ---
        def add_section(title: str, content: str):
            if not content:
                return
            h = doc.add_paragraph()
            run = h.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            h.paragraph_format.space_before = Pt(12)

            # Linha separadora
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_before = Pt(0)
            p_sep.paragraph_format.space_after = Pt(4)
            run_sep = p_sep.add_run('─' * 60)
            run_sep.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)
            run_sep.font.size = Pt(8)

            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(8)

        add_section('Pauta', ata.pauta)
        add_section('Deliberações', ata.deliberacoes)

        # --- Tabela de ações ---
        acoes = ata.acoes or []
        if acoes:
            h = doc.add_paragraph()
            run = h.add_run('PLANO DE AÇÃO')
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            h.paragraph_format.space_before = Pt(12)

            table = doc.add_table(rows=1 + len(acoes), cols=3)
            table.style = 'Table Grid'

            headers = ['Ação', 'Responsável', 'Prazo']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                _set_cell_bg(cell, '4F46E5')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for row_idx, acao in enumerate(acoes, start=1):
                row = table.rows[row_idx]
                row.cells[0].text = acao.get('acao', '')
                row.cells[1].text = acao.get('responsavel', '')
                row.cells[2].text = acao.get('prazo', '')
                if row_idx % 2 == 0:
                    for cell in row.cells:
                        _set_cell_bg(cell, 'EEF2FF')

        # --- Conteúdo completo (opcional — seção colapsada) ---
        doc.add_paragraph()
        add_section('Registro Completo da Transcrição', ata.conteudo_completo)

        # --- Rodapé ---
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(f'Documento gerado automaticamente via GIS Plataforma')
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
